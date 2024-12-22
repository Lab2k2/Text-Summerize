import os
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_core.prompts import PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from docx import Document as DocxDocument   # Đảm bảo cài đặt thư viện python-docx
from llm_model import llm

# Thiết lập cấu hình Streamlit
st.set_page_config(page_title="Tóm Tắt Văn Bản", layout="wide")

# Thiết lập User-Agent
USER_AGENT = "MyApp/1.0"
HEADERS = {"User-Agent": USER_AGENT}

# Prompt tóm tắt
map_prompt = """
Bạn là một chuyên gia tóm tắt. Hãy tóm tắt văn bản sau một cách ngắn gọn bằng tiếng Việt:

{text}
"""
map_template = PromptTemplate(template=map_prompt, input_variables=["text"])

# Hàm đọc tệp .txt
def read_txt(file):
    return file.read().decode("utf-8")

# Hàm đọc tệp .docx
def read_docx(file):
    doc = DocxDocument(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# Tạo map-reduce chain
chain = load_summarize_chain(llm, chain_type="map_reduce", map_prompt=map_template)

# Layout Streamlit: Hai cột ngang
col1, col2 = st.columns(2)

# Nửa bên trái: Chọn tệp và hiển thị nội dung
with col1:
    st.header("Chọn Tệp Tin")
    uploaded_file = st.file_uploader("", type=["txt", "docx"])
    
    if uploaded_file is not None:
        # Đọc tệp
        if uploaded_file.name.endswith(".txt"):
            text_content = read_txt(uploaded_file)
        elif uploaded_file.name.endswith(".docx"):
            text_content = read_docx(uploaded_file)

        # Hiển thị nội dung tệp
        st.subheader("Nội Dung Tệp")
        st.text_area("", text_content, height=400)

        # Chuyển đổi nội dung thành Document
        documents = [Document(page_content=text_content)]

        # Chia tài liệu thành các phần nhỏ hơn
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        split_docs = text_splitter.split_documents(documents)

# Nửa bên phải: Tóm tắt nội dung
with col2:
    st.header("Tóm Tắt Nội Dung")
    if uploaded_file is not None:
        if st.button("Tóm Tắt"):
            try:
                # Tóm tắt nội dung
                final_summary = chain.invoke({"input_documents": split_docs})

                # Kiểm tra và hiển thị kết quả tóm tắt từ dict
                if isinstance(final_summary, dict) and 'output_text' in final_summary:
                    st.write(final_summary['output_text']) 
                else:
                    st.error("Không có kết quả tóm tắt.")
            except Exception as e:
                st.error(f"Error during summarization: {e}")
    else:
        st.write("Hãy chọn tệp tin để tóm tắt.")
